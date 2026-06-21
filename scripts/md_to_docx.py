#!/usr/bin/env python
"""Convert a Markdown file to a formatted .docx (headings, tables, code, lists).

    .venv/bin/python scripts/md_to_docx.py grading-rubric-proposal.md

Writes <name>.docx next to the input. Uses markdown -> HTML -> python-docx.
"""
from __future__ import annotations

import sys
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
CODE_BG = "F2F2F2"


def add_runs(paragraph, node, mono=False, bold=False, italic=False):
    """Recursively add styled runs for an inline HTML node."""
    if isinstance(node, NavigableString):
        r = paragraph.add_run(str(node))
        r.bold = bold or None
        r.italic = italic or None
        if mono:
            r.font.name = "Consolas"; r.font.size = Pt(9.5)
        return
    name = node.name
    b = bold or name in ("strong", "b")
    i = italic or name in ("em", "i")
    m = mono or name == "code"
    for child in node.children:
        add_runs(paragraph, child, mono=m, bold=b, italic=i)


def shade(cell_or_para_xml, fill):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill)
    cell_or_para_xml.append(shd)


def convert(md_path: Path) -> Path:
    html = markdown.markdown(md_path.read_text(),
                             extensions=["tables", "fenced_code", "sane_lists"])
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    for el in soup.find_all(recursive=False) or soup.children:
        if isinstance(el, NavigableString):
            continue
        tag = el.name
        if tag in ("h1", "h2", "h3", "h4"):
            level = int(tag[1])
            h = doc.add_heading(level=level if level <= 4 else 4)
            add_runs(h, el)
            for run in h.runs:
                run.font.color.rgb = ACCENT
        elif tag == "p":
            p = doc.add_paragraph()
            add_runs(p, el)
        elif tag in ("ul", "ol"):
            style = "List Bullet" if tag == "ul" else "List Number"
            for li in el.find_all("li", recursive=False):
                p = doc.add_paragraph(style=style)
                add_runs(p, li)
        elif tag == "blockquote":
            for sub in el.find_all(["p"], recursive=False) or [el]:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                add_runs(p, sub, italic=True)
        elif tag == "pre":
            code = el.get_text()
            p = doc.add_paragraph()
            r = p.add_run(code.rstrip("\n"))
            r.font.name = "Consolas"; r.font.size = Pt(9)
            shade(p._p.get_or_add_pPr(), CODE_BG)
        elif tag == "table":
            rows = el.find_all("tr")
            if not rows:
                continue
            ncol = max(len(r.find_all(["td", "th"])) for r in rows)
            t = doc.add_table(rows=0, cols=ncol)
            t.style = "Light Grid Accent 1"
            for r in rows:
                cells = r.find_all(["td", "th"])
                row_cells = t.add_row().cells
                for ci, c in enumerate(cells):
                    cell = row_cells[ci]
                    cell.paragraphs[0].text = ""
                    add_runs(cell.paragraphs[0], c, bold=(c.name == "th"))
        elif tag == "hr":
            doc.add_paragraph("_" * 40)

    out = md_path.with_suffix(".docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print(__doc__); sys.exit(1)
    out = convert(Path(sys.argv[1]))
    print(f"wrote {out}")
